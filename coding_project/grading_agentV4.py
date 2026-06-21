"""
grading_agentV4.py  —  Automated Grading Agent
Cisco Cyber Operations Grading System

Changes from V3:
  - PDF pages converted to images via pdf2image before sending to API
  - Pre-grading verification step counts images in each PDF
  - Grading stops and file moves to failed/ folder if image count too low
  - Hardcoded poppler path — no Windows PATH variable needed
  - Each page sent as a separate image in the API message content
  - More reliable screenshot detection for Word-generated PDFs

Usage:
    python grading_agentV4.py

Dependencies:
    pip install anthropic python-dotenv python-docx pdf2image
    poppler installed at: C:\\poppler\\poppler-26.02.0\\Library\\bin
"""

import os
import sys
import csv
import time
import shutil
import base64
import re
import threading
from datetime import datetime
from pathlib import Path
from io import BytesIO

import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from pdf2image import convert_from_path
from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError

# ── Load API key ──────────────────────────────────────────────────────────────
load_dotenv(Path("I:/Visual_Studio_Code/.env"))
API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not API_KEY:
    print("\n  ERROR: ANTHROPIC_API_KEY not found in .env file.")
    print("  Check I:\\Visual_Studio_Code\\.env has your API key.\n")
    sys.exit(1)

# ── Configuration ─────────────────────────────────────────────────────────────
BASE_DIR       = Path("I:/Visual_Studio_Code")
POLL_SECS      = 30
MODEL          = "claude-opus-4-6"
MAX_TOKENS     = 4096
AGENT_VERSION  = "V4"
MAX_RETRIES    = 3
POPPLER_PATH   = r"C:\poppler\poppler-26.02.0\Library\bin"

# Minimum number of pages expected in a student submission
# If PDF renders fewer pages than this it is likely corrupted or unreadable
MIN_PAGES      = 3

# DPI for page rendering — higher is clearer but slower and larger
# 150 is a good balance for screenshot readability
RENDER_DPI     = 150
# ──────────────────────────────────────────────────────────────────────────────


def pick_course_folder() -> Path:
    """Scan BASE_DIR for course subfolders and prompt instructor to pick one."""
    course_folders = sorted([
        f for f in BASE_DIR.iterdir()
        if f.is_dir() and (f / "roster.csv").exists()
    ])

    if not course_folders:
        print(f"\n  ERROR: No course folders found under {BASE_DIR}")
        sys.exit(1)

    print("\n  ── Course Folders Found ────────────────────────────")
    for i, folder in enumerate(course_folders, start=1):
        print(f"  {i:>3}.  {folder.name}")
    print("  ────────────────────────────────────────────────────")

    while True:
        raw = input("\n  Enter course number: ").strip()
        if raw.isdigit():
            num = int(raw)
            if 1 <= num <= len(course_folders):
                selected = course_folders[num - 1]
                confirm  = input(
                    f"  Confirm: {selected.name}  (y/n): "
                ).strip().lower()
                if confirm == "y":
                    return selected
                print("  Not confirmed.  Please re-enter.")
                continue
        print(f"  Invalid.  Enter a number between 1 and {len(course_folders)}.")


def load_roster(course_dir: Path) -> dict:
    """Load roster as initials -> full name lookup."""
    roster_file = course_dir / "roster.csv"
    lookup = {}
    with open(roster_file, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        reader.fieldnames = [h.strip().lower() for h in reader.fieldnames]
        for row in reader:
            clean    = {k.strip().lower(): v for k, v in row.items()}
            initials = clean["initials"].strip().upper()
            lookup[initials] = (
                f"{clean['first_name'].strip()} {clean['last_name'].strip()}"
            )
    return lookup


def parse_queue_filename(filename: str):
    """
    Extract initials, lab number, and regrade version from queue filename.

    Handles:
        AF_lab04_pending.pdf        -> (AF, 4, None)
        AF_lab04_pending_r2.pdf     -> (AF, 4, r2)
        AF_lab04_pending_r3.pdf     -> (AF, 4, r3)
    """
    stem          = Path(filename).stem
    regrade_match = re.search(r'_r(\d+)$', stem)
    regrade_ver   = f"r{regrade_match.group(1)}" if regrade_match else None
    match         = re.match(r'^([A-Z0-9]+)_lab(\d+)_pending', stem, re.IGNORECASE)
    if match:
        return match.group(1).upper(), int(match.group(2)), regrade_ver
    return None, None, None


def find_template_files(course_dir: Path, lab_num: int):
    """
    Locate answer template and grading notes for a given lab number.
    Naming convention from lead instructor:
        2026SU_Lab_#_-_ILM.docx
        Lab_#_Grading_Notes.txt
    """
    templates_dir = course_dir / "templates"
    answers_file  = templates_dir / f"2026SU_Lab_{lab_num}_-_ILM.docx"
    rubric_file   = templates_dir / f"Lab_{lab_num}_Grading_Notes.txt"
    answers       = answers_file if answers_file.exists() else None
    rubric        = rubric_file  if rubric_file.exists()  else None
    return answers, rubric


def read_text_file(file_path: Path) -> str:
    """Read a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_docx_text(file_path: Path) -> str:
    """Extract plain text from a docx file."""
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        return f"[Could not extract docx text: {e}]"


def pdf_to_images(pdf_path: Path) -> list:
    """
    Convert PDF pages to PIL images using pdf2image and poppler.
    Returns list of PIL Image objects, one per page.
    Raises exception if conversion fails.
    """
    return convert_from_path(
        str(pdf_path),
        dpi=RENDER_DPI,
        poppler_path=POPPLER_PATH,
        fmt="jpeg",
        thread_count=2,
    )


def image_to_base64(pil_image) -> str:
    """Convert a PIL image to base64 string for API submission."""
    buffer = BytesIO()
    pil_image.save(buffer, format="JPEG", quality=85)
    return base64.standard_b64encode(buffer.getvalue()).decode("utf-8")


def verify_pdf_renders(pdf_path: Path, failed_dir: Path) -> tuple[bool, list]:
    """
    Verify PDF renders correctly and returns enough pages.
    Returns (success, pages) where pages is list of PIL images.
    Moves file to failed/ folder if verification fails.
    """
    print(f"  Verifying PDF renders correctly...")

    try:
        pages = pdf_to_images(pdf_path)
    except PDFInfoNotInstalledError:
        print(f"  ERROR: Poppler not found at {POPPLER_PATH}")
        print(f"  Check poppler installation.")
        return False, []
    except PDFPageCountError:
        print(f"  ERROR: Cannot read page count — PDF may be corrupted.")
        _move_to_failed(pdf_path, failed_dir, "corrupted_pdf")
        return False, []
    except Exception as e:
        print(f"  ERROR: PDF conversion failed — {e}")
        _move_to_failed(pdf_path, failed_dir, "conversion_failed")
        return False, []

    page_count = len(pages)
    print(f"  Pages rendered: {page_count}")

    if page_count < MIN_PAGES:
        print(f"  ERROR: Only {page_count} page(s) rendered.")
        print(f"  Minimum expected: {MIN_PAGES}")
        print(f"  PDF may be corrupted or nearly empty.")
        _move_to_failed(pdf_path, failed_dir, "too_few_pages")
        return False, []

    print(f"  ✓  PDF verification passed — {page_count} pages ready for grading.")
    return True, pages


def _move_to_failed(pdf_path: Path, failed_dir: Path, reason: str):
    """Move a problem PDF to the failed/ folder with reason suffix."""
    failed_dir.mkdir(exist_ok=True)
    dest = failed_dir / f"{pdf_path.stem}_{reason}{pdf_path.suffix}"
    shutil.move(str(pdf_path), str(dest))
    print(f"  Moved to failed/ as: {dest.name}")
    print(f"  Check this PDF manually and re-intake if needed.")


def build_grading_prompt(answers_text: str, rubric_text: str,
                         student_name: str, lab_num: int,
                         page_count: int) -> str:
    """Build the system prompt for the grading API call."""
    return f"""You are an experienced Cisco Cyber Operations instructor grading a student lab submission.

STUDENT: {student_name}
LAB: {lab_num}

You have been provided:
1. The ANSWER TEMPLATE — the correct answers and required screenshots for this lab
2. The GRADING NOTES — the lead instructor rubric with specific deductions
3. The STUDENT SUBMISSION — {page_count} pages rendered as images from the student PDF
   Each page image shows exactly what the student submitted including all screenshots.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
IMPORTANT — HOW TO READ THE SUBMISSION
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

The student submission is rendered as page images.
You can see ALL content including screenshots pasted by the student.
Read each page carefully — text answers AND screenshots are both visible.
Do not assume a screenshot is missing unless you have checked every page.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
SCREENSHOT GRADING RULES
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

TYPE 1 \u2014 SCREENSHOT IS THE ANSWER
  The template asks for a screenshot with no qualifying statement below it.
  Grade the screenshot directly against the answer template.
  If missing or clearly incorrect \u2014 deduct full points.
  If present but unclear \u2014 deduct partial points and explain.

TYPE 2 \u2014 SCREENSHOT IS REQUIRED BUT NOT THE ANSWER
  The template contains this statement near the question:
  "Your screenshot is not used to answer the questions below,
  but you must provide a correct screenshot on the previous
  page, or your answers are not accepted"

  STEP A \u2014 Verify screenshot exists on the previous page (pass/fail)
    If MISSING: flag as separate deduction, still grade text answers,
    warn that instructor may not accept answers without screenshot.
    If PRESENT: do not use screenshot to evaluate text answers.

  STEP B \u2014 Grade TEXT ANSWERS below the statement independently.

SCREENSHOT FORMATTING DEDUCTIONS (max 11 points total):
  - Screenshot too small to read output clearly
  - Wrong area of screen captured
  - NDG username not visible
  - Multiple screenshots where one was required

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
GRADING NOTES USAGE
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

Use grading notes as PRIMARY source for deductions.
Expand and improve the language \u2014 do not copy word for word.
Connect errors to real Cisco Cyber Ops or SOC security concepts.
Apply the most appropriate deduction level to what was submitted.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
OUTPUT FORMAT \u2014 USE EXACTLY THIS STRUCTURE
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

EXECUTIVE_SUMMARY_START
Total Score: [X]/[total]
Screenshot Formatting Deduction: [X] points
Missing Required Screenshots: [X] questions affected

Q: [Part X Step Xx] | Score: [X]/[X] | [Correct  OR  brief reason]
[continue for all questions]
EXECUTIVE_SUMMARY_END

DETAILED_FEEDBACK_START
[Only include questions where points were deducted]

QUESTION: [Part X Step Xx] \u2014 [X]/[X] points
SCREENSHOT_STATUS: [Not applicable / Present / Missing / Present but unclear]
WHAT_YOU_SUBMITTED: [brief summary of what the student wrote or showed]
WHAT_WAS_WRONG: [specific explanation of the error]
WHY_IT_MATTERS: [connect to Cisco Cyber Ops or SOC security concept]
HOW_TO_FIX: [clear corrective guidance for resubmission]

[repeat for each incorrect question]
DETAILED_FEEDBACK_END

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
GRADING MATERIALS
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

ANSWER TEMPLATE:
{answers_text}

GRADING NOTES:
{rubric_text}
"""


def call_grading_api(client: anthropic.Anthropic,
                     prompt: str,
                     pages: list,
                     student_file: Path) -> str:
    """
    Send rendered PDF page images to Claude API.
    Each page is sent as a separate image in the message content.
    Falls back to text extraction for DOCX files.
    """
    ext = student_file.suffix.lower()

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            start     = time.time()
            stop_flag = threading.Event()

            def show_progress():
                while not stop_flag.is_set():
                    elapsed = int(time.time() - start)
                    print(
                        f"\r  [{datetime.now().strftime('%H:%M:%S')}]  "
                        f"Waiting for API response...  ({elapsed}s)",
                        end="", flush=True
                    )
                    time.sleep(5)

            t = threading.Thread(target=show_progress, daemon=True)
            t.start()

            if ext == ".pdf" and pages:
                # Build content list — one image block per page
                content = []
                for i, page in enumerate(pages):
                    content.append({
                        "type": "image",
                        "source": {
                            "type":       "base64",
                            "media_type": "image/jpeg",
                            "data":       image_to_base64(page),
                        }
                    })
                    # Add page label as text between pages for context
                    content.append({
                        "type": "text",
                        "text": f"[Page {i + 1} of {len(pages)}]"
                    })

                content.append({
                    "type": "text",
                    "text": (
                        "The pages above are the complete student submission. "
                        "Please grade this submission according to the rubric "
                        "and answer template provided in your instructions."
                    )
                })

                message = client.messages.create(
                    model=MODEL,
                    max_tokens=MAX_TOKENS,
                    system=prompt,
                    messages=[{"role": "user", "content": content}],
                )

            else:
                # DOCX fallback — extract text
                student_text = extract_docx_text(student_file)
                message      = client.messages.create(
                    model=MODEL,
                    max_tokens=MAX_TOKENS,
                    system=prompt,
                    messages=[
                        {
                            "role": "user",
                            "content": (
                                "Please grade this student submission according "
                                "to the rubric and answer template provided.\n\n"
                                f"STUDENT SUBMISSION:\n{student_text}"
                            )
                        }
                    ],
                )

            stop_flag.set()
            t.join()
            elapsed = int(time.time() - start)
            print(
                f"\r  [{datetime.now().strftime('%H:%M:%S')}]  "
                f"Response received.  ({elapsed}s)          "
            )
            return message.content[0].text

        except Exception as e:
            stop_flag.set()
            t.join()
            print(f"\n  Attempt {attempt}/{MAX_RETRIES} failed: {e}")
            if attempt < MAX_RETRIES:
                wait = 15 * attempt
                print(f"  Retrying in {wait} seconds...")
                time.sleep(wait)
            else:
                raise


def parse_grade_response(response: str) -> tuple[str, str, int, int]:
    """Parse Claude response into summary, detailed, score, total."""
    summary_match = re.search(
        r'EXECUTIVE_SUMMARY_START(.*?)EXECUTIVE_SUMMARY_END',
        response, re.DOTALL
    )
    summary_text  = summary_match.group(1).strip() if summary_match else response

    detailed_match = re.search(
        r'DETAILED_FEEDBACK_START(.*?)DETAILED_FEEDBACK_END',
        response, re.DOTALL
    )
    detailed_text  = detailed_match.group(1).strip() if detailed_match else ""

    score_match = re.search(r'Total Score:\s*(\d+)/(\d+)', summary_text)
    score = int(score_match.group(1)) if score_match else 0
    total = int(score_match.group(2)) if score_match else 96

    return summary_text, detailed_text, score, total


def write_feedback_docx(course_dir: Path, student_name: str, initials: str,
                        lab_num: int, summary_text: str, detailed_text: str,
                        score: int, total: int,
                        regrade_ver: str | None) -> Path:
    """Write structured feedback docx. Versioned if regraded."""
    grades_dir  = course_dir / "grades" / f"lab{lab_num:02d}"
    grades_dir.mkdir(parents=True, exist_ok=True)
    ver_suffix  = f"_{regrade_ver}" if regrade_ver else ""
    output_path = grades_dir / f"{initials}_lab{lab_num:02d}_feedback{ver_suffix}.docx"
    is_regrade  = regrade_ver is not None
    doc         = Document()

    section               = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    def add_heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)
        return p

    def add_para(text, bold=False, size=11, color=None, indent=0):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    # ── Agent version stamp ───────────────────────────────────────
    stamp = doc.add_paragraph()
    run   = stamp.add_run(
        f"Grading Agent: {AGENT_VERSION}  |  "
        f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"{'Regraded: ' + regrade_ver.upper() if is_regrade else 'Original Grade'}"
    )
    run.font.size      = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic         = True

    doc.add_paragraph()

    # ── SOC Standard Notice ───────────────────────────────────────
    notice     = doc.add_paragraph()
    notice_run = notice.add_run(
        "GRADING NOTE \u2014 PROFESSIONAL SOC STANDARD\n\n"
        "This lab has been graded to the documentation standard required of a "
        "Level 2 Security Operations Center analyst. In a real SOC environment, "
        "incomplete commands, unclear screenshots, or imprecise descriptions in "
        "an incident report can cause a security alert to be misrouted, delayed, "
        "or missed entirely \u2014 putting an organization at risk.\n\n"
        "To recognize that you are still developing these professional skills, "
        "20 points have been added to your raw score. Your final grade reflects "
        "both your current progress and the high standard this career demands.\n\n"
        "Please review every item in the detailed feedback section below \u2014 "
        "even if you choose not to resubmit. Understanding exactly what was wrong "
        "and why it matters will make you a stronger analyst. These details are "
        "what separates a good SOC report from a great one."
    )
    notice_run.font.size      = Pt(10)
    notice_run.font.color.rgb = RGBColor(0, 70, 127)

    doc.add_paragraph()

    # ── Header ────────────────────────────────────────────────────
    title = "STUDENT FEEDBACK REPORT"
    if is_regrade:
        title += f"  ({regrade_ver.upper()})"
    add_heading(title, level=1, color=(0, 70, 127))
    doc.add_paragraph()

    table       = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Student:", student_name),
        ("Lab:",     f"Lab {lab_num}"),
        ("Date:",    datetime.now().strftime("%Y-%m-%d")),
    ]):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────
    add_heading("EXECUTIVE SUMMARY", level=2, color=(0, 70, 127))

    score_pct   = round((score / total) * 100) if total > 0 else 0
    adjusted    = min(score + 20, total)
    score_color = (0, 128, 0) if score_pct >= 70 else (200, 0, 0)

    add_para(f"Raw Score:       {score} / {total}", bold=True, size=12)
    add_para(
        f"Adjusted Score:  {adjusted} / {total}  ({min(score_pct + 20, 100)}%)",
        bold=True, size=13, color=score_color
    )
    doc.add_paragraph()

    if "missing required screenshot" in summary_text.lower():
        add_para(
            "\u26a0  One or more required screenshots were missing. "
            "See detailed feedback. Text answer points are preserved "
            "but your instructor may not accept answers without "
            "screenshots per lab policy.",
            bold=True, size=10, color=(180, 80, 0)
        )
        doc.add_paragraph()

    for line in summary_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("Total Score:"):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        if line.startswith("Q:"):
            run = p.add_run(line)
            run.font.color.rgb = (
                RGBColor(0, 128, 0) if "Correct" in line
                else RGBColor(180, 0, 0)
            )
        else:
            p.add_run(line)

    doc.add_paragraph()

    # ── Detailed Feedback ─────────────────────────────────────────
    if detailed_text.strip():
        add_heading("DETAILED FEEDBACK", level=2, color=(0, 70, 127))
        add_para(
            "The following questions had deductions. "
            "Review each item carefully before resubmitting.",
            size=10
        )
        doc.add_paragraph()

        for block in re.split(r'(?=QUESTION:)', detailed_text):
            block = block.strip()
            if not block:
                continue
            for line in block.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                if line.startswith("QUESTION:"):
                    p = doc.add_heading(
                        line.replace("QUESTION:", "").strip(), level=3
                    )
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(180, 0, 0)
                elif line.startswith("SCREENSHOT_STATUS:"):
                    status = line.replace("SCREENSHOT_STATUS:", "").strip()
                    color  = (180, 80, 0) if "Missing" in status else (100, 100, 100)
                    add_para(
                        f"Screenshot Status:  {status}",
                        bold=True, size=10, color=color, indent=0.25
                    )
                elif line.startswith("WHAT_YOU_SUBMITTED:"):
                    add_para("What You Submitted:", bold=True)
                    add_para(line.replace("WHAT_YOU_SUBMITTED:", "").strip(), indent=0.25)
                elif line.startswith("WHAT_WAS_WRONG:"):
                    add_para("What Was Wrong:", bold=True)
                    add_para(line.replace("WHAT_WAS_WRONG:", "").strip(), indent=0.25)
                elif line.startswith("WHY_IT_MATTERS:"):
                    add_para("Why It Matters:", bold=True)
                    add_para(line.replace("WHY_IT_MATTERS:", "").strip(), indent=0.25)
                elif line.startswith("HOW_TO_FIX:"):
                    add_para("How To Fix It:", bold=True)
                    add_para(line.replace("HOW_TO_FIX:", "").strip(), indent=0.25)
                else:
                    add_para(line)
            doc.add_paragraph("\u2500" * 60)

    doc.save(str(output_path))
    return output_path


def write_grade_log(course_dir: Path, initials: str, student_name: str,
                    lab_num: int, score: int, total: int,
                    feedback_file: str, regrade_ver: str | None,
                    page_count: int):
    """Append grade record to grade_log.csv."""
    log_file    = course_dir / "grade_log.csv"
    file_exists = log_file.exists()
    with open(log_file, "a", newline="", encoding="utf-8") as f:
        fieldnames = [
            "timestamp", "initials", "student_name", "lab",
            "score", "adjusted_score", "total", "percent",
            "regrade", "agent_version", "pages_rendered", "feedback_file"
        ]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        if not file_exists:
            writer.writeheader()
        pct      = round((score / total) * 100) if total > 0 else 0
        adjusted = min(score + 20, total)
        writer.writerow({
            "timestamp":      datetime.now().strftime("%Y-%m-%d %H:%M"),
            "initials":       initials,
            "student_name":   student_name,
            "lab":            lab_num,
            "score":          score,
            "adjusted_score": adjusted,
            "total":          total,
            "percent":        f"{min(pct + 20, 100)}%",
            "regrade":        regrade_ver if regrade_ver else "No",
            "agent_version":  AGENT_VERSION,
            "pages_rendered": page_count,
            "feedback_file":  feedback_file,
        })


def process_submission(client: anthropic.Anthropic, course_dir: Path,
                       roster: dict, file: Path):
    """Grade one student submission file."""
    print(f"\n  {'='*54}")
    print(f"  Grading:  {file.name}")
    print(f"  {'='*54}")

    initials, lab_num, regrade_ver = parse_queue_filename(file.name)
    if not initials or not lab_num:
        print(f"  ERROR: Cannot parse filename {file.name} — skipping.")
        return

    student_name  = roster.get(initials, initials)
    regrade_label = f"  (Regrade {regrade_ver.upper()})" if regrade_ver else ""
    print(f"  Student:  {student_name}{regrade_label}")
    print(f"  Lab:      {lab_num}")

    answers_path, rubric_path = find_template_files(course_dir, lab_num)
    if not answers_path:
        print(f"  ERROR: Answer template not found for Lab {lab_num}")
        print(f"  Expected: templates/2026SU_Lab_{lab_num}_-_ILM.docx")
        return
    if not rubric_path:
        print(f"  ERROR: Grading notes not found for Lab {lab_num}")
        print(f"  Expected: templates/Lab_{lab_num}_Grading_Notes.txt")
        return

    print(f"  Template: {answers_path.name}")
    print(f"  Rubric:   {rubric_path.name}")

    # Verify and render PDF pages
    pages      = []
    page_count = 0
    if file.suffix.lower() == ".pdf":
        failed_dir = course_dir / "failed"
        ok, pages  = verify_pdf_renders(file, failed_dir)
        if not ok:
            return
        page_count = len(pages)

    answers_text = extract_docx_text(answers_path)
    rubric_text  = read_text_file(rubric_path)
    prompt       = build_grading_prompt(
        answers_text, rubric_text, student_name, lab_num, page_count
    )

    print(f"  [{datetime.now().strftime('%H:%M:%S')}]  "
          f"Sending {page_count} pages to Claude API...")
    try:
        response = call_grading_api(client, prompt, pages, file)
    except Exception as e:
        print(f"\n  ERROR: API call failed after {MAX_RETRIES} attempts — {e}")
        return

    summary_text, detailed_text, score, total = parse_grade_response(response)
    adjusted = min(score + 20, total)
    print(f"  Raw Score:      {score}/{total}")
    print(f"  Adjusted Score: {adjusted}/{total}")

    feedback_path = write_feedback_docx(
        course_dir, student_name, initials, lab_num,
        summary_text, detailed_text, score, total, regrade_ver
    )
    print(f"  Feedback: {feedback_path.name}")

    processed_dir = course_dir / "processed"
    processed_dir.mkdir(exist_ok=True)
    dest = processed_dir / file.name
    if dest.exists():
        ts   = datetime.now().strftime("%H%M%S")
        dest = processed_dir / f"{file.stem}_{ts}{file.suffix}"
    shutil.move(str(file), str(dest))

    write_grade_log(
        course_dir, initials, student_name, lab_num,
        score, total, feedback_path.name, regrade_ver, page_count
    )

    print(f"  ✓  Complete — moved to processed/")


def main():
    print("\n" + "="*54)
    print(f"  Cisco Cyber Ops  \u2014  Grading Agent {AGENT_VERSION}")
    print("="*54)

    course_dir = pick_course_folder()
    print(f"\n  Working in:  {course_dir.name}")

    roster    = load_roster(course_dir)
    client    = anthropic.Anthropic(api_key=API_KEY)
    queue_dir = course_dir / "queue"

    print(f"\n  Watching queue every {POLL_SECS} seconds.")
    print(f"  Press Ctrl+C to stop.\n")

    try:
        while True:
            pending = sorted([
                f for f in queue_dir.iterdir()
                if f.is_file()
                and "_pending" in f.name.lower()
                and f.suffix.lower() in {".pdf", ".docx"}
            ])

            if pending:
                print(f"  Found {len(pending)} file(s) in queue.")
                for file in pending:
                    process_submission(client, course_dir, roster, file)
            else:
                print(
                    f"  [{datetime.now().strftime('%H:%M:%S')}]"
                    f"  Queue empty \u2014 waiting..."
                )
            time.sleep(POLL_SECS)

    except KeyboardInterrupt:
        print(f"\n\n  Grading agent {AGENT_VERSION} stopped.  Grade log saved.\n")


if __name__ == "__main__":
    main()
