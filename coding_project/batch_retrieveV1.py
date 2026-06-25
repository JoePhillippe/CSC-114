"""
batch_retrieveV1.py  —  Batch Results Retrieval Script
Cisco Cyber Operations Grading System

Usage:
    python batch_retrieveV1.py

Workflow:
    1. Prompts you to select which course folder to work with
    2. Reads batch_jobs.csv to find pending batch job IDs
    3. Checks Anthropic API for batch status
    4. If complete — downloads results and writes feedback docx files
    5. Updates batch_jobs.csv status to completed
    6. Logs grades to grade_log.csv

Run this the day after submitting with grading_agentV5.py.
Can be run multiple times safely — already processed results are skipped.

Dependencies:
    pip install anthropic python-dotenv python-docx
"""

import os
import sys
import csv
import re
from datetime import datetime
from pathlib import Path

import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# ── Load API key ──────────────────────────────────────────────────────────────
load_dotenv(Path("I:/Visual_Studio_Code/.env"))
API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not API_KEY:
    print("\n  ERROR: ANTHROPIC_API_KEY not found in .env file.")
    sys.exit(1)

# ── Configuration ─────────────────────────────────────────────────────────────
BASE_DIR      = Path("I:/Visual_Studio_Code")
AGENT_VERSION = "V5"
# ──────────────────────────────────────────────────────────────────────────────


def pick_course_folder() -> Path:
    course_folders = sorted([
        f for f in BASE_DIR.iterdir()
        if f.is_dir() and (f / "roster.csv").exists()
    ])
    if not course_folders:
        print(f"\n  ERROR: No course folders found under {BASE_DIR}")
        sys.exit(1)
    print("\n  ── Course Folders Found ────────────────────────────")
    for i, folder in enumerate(course_folders, start=1):
        print(f"  {i:>3}.  {folder.name}")
    print("  ────────────────────────────────────────────────────")
    while True:
        raw = input("\n  Enter course number: ").strip()
        if raw.isdigit():
            num = int(raw)
            if 1 <= num <= len(course_folders):
                selected = course_folders[num - 1]
                confirm  = input(f"  Confirm: {selected.name}  (y/n): ").strip().lower()
                if confirm == "y":
                    return selected
                print("  Not confirmed.  Please re-enter.")
                continue
        print(f"  Invalid.  Enter a number between 1 and {len(course_folders)}.")


def load_batch_manifest(course_dir: Path) -> list[dict]:
    """Load pending batch jobs from batch_jobs.csv."""
    manifest_file = course_dir / "batch_jobs.csv"
    if not manifest_file.exists():
        return []
    jobs = []
    with open(manifest_file, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row.get("status") == "pending":
                jobs.append(dict(row))
    return jobs


def update_manifest_status(course_dir: Path, custom_id: str, status: str):
    """Update a single record status in batch_jobs.csv."""
    manifest_file = course_dir / "batch_jobs.csv"
    rows = []
    with open(manifest_file, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames
        for row in reader:
            if row["custom_id"] == custom_id and row["status"] == "pending":
                row["status"] = status
            rows.append(row)
    with open(manifest_file, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def parse_grade_response(response: str) -> tuple[str, str, int, int]:
    """Parse Claude response into summary, detailed, score, total."""
    summary_match = re.search(
        r'EXECUTIVE_SUMMARY_START(.*?)EXECUTIVE_SUMMARY_END',
        response, re.DOTALL
    )
    summary_text  = summary_match.group(1).strip() if summary_match else response

    detailed_match = re.search(
        r'DETAILED_FEEDBACK_START(.*?)DETAILED_FEEDBACK_END',
        response, re.DOTALL
    )
    detailed_text  = detailed_match.group(1).strip() if detailed_match else ""

    # Method 1 — sum Q: matrix lines
    q_matches = re.findall(r'Q:.*?\|\s*(\d+)/(\d+)\s*\|', summary_text)
    if q_matches:
        score_summed = sum(int(e) for e, p in q_matches)
        total_summed = sum(int(p) for e, p in q_matches)
    else:
        score_summed = None
        total_summed = None

    # Method 2 — read Total Score line
    score_match  = re.search(r'Total Score:\s*(\d+)/(\d+)', summary_text)
    score_direct = int(score_match.group(1)) if score_match else None
    total_direct = int(score_match.group(2)) if score_match else None

    # Method 3 — Points Possible line
    possible_match = re.search(r'Points Possible:\s*(\d+)', summary_text)
    total_possible = int(possible_match.group(1)) if possible_match else None

    # Best score — prefer summed Q lines
    if score_summed is not None and score_summed > 0:
        score = score_summed
    elif score_direct is not None:
        score = score_direct
    else:
        score = 0

    # Best total — prefer Points Possible
    if total_possible is not None:
        total = total_possible
    elif total_summed is not None and total_summed > 0:
        total = total_summed
    elif total_direct is not None:
        total = total_direct
    else:
        total = 96

    # Apply screenshot formatting deduction
    screenshot_match = re.search(
        r'Screenshot Formatting Deduction:\s*(\d+)', summary_text
    )
    if screenshot_match:
        score = max(0, score - int(screenshot_match.group(1)))

    return summary_text, detailed_text, score, total


def write_feedback_docx(course_dir: Path, student_name: str, initials: str,
                        lab_num: int, summary_text: str, detailed_text: str,
                        score: int, total: int,
                        regrade_ver: str | None) -> Path:
    """Write structured feedback docx to grades/lab##/ folder."""
    grades_dir  = course_dir / "grades" / f"lab{lab_num:02d}"
    grades_dir.mkdir(parents=True, exist_ok=True)
    ver_suffix  = f"_{regrade_ver}" if regrade_ver else ""
    output_path = grades_dir / f"{initials}_lab{lab_num:02d}_feedback{ver_suffix}.docx"
    is_regrade  = regrade_ver is not None
    doc         = Document()

    section               = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    def add_heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)
        return p

    def add_para(text, bold=False, size=11, color=None, indent=0):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    # ── Agent version stamp ───────────────────────────────────────
    stamp = doc.add_paragraph()
    run   = stamp.add_run(
        f"Grading Agent: {AGENT_VERSION}  |  "
        f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"{'Regraded: ' + regrade_ver.upper() if is_regrade else 'Original Grade'}"
    )
    run.font.size      = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic         = True

    doc.add_paragraph()

    # ── SOC Standard Notice ───────────────────────────────────────
    notice     = doc.add_paragraph()
    notice_run = notice.add_run(
        "GRADING NOTE \u2014 PROFESSIONAL SOC STANDARD\n\n"
        "This lab has been graded to the documentation standard required of a "
        "Level 2 Security Operations Center analyst. In a real SOC environment, "
        "incomplete commands, unclear screenshots, or imprecise descriptions in "
        "an incident report can cause a security alert to be misrouted, delayed, "
        "or missed entirely \u2014 putting an organization at risk.\n\n"
        "To recognize that you are still developing these professional skills, "
        "20 points have been added to your raw score. Your final grade reflects "
        "both your current progress and the high standard this career demands.\n\n"
        "Please review every item in the detailed feedback section below \u2014 "
        "even if you choose not to resubmit. Understanding exactly what was wrong "
        "and why it matters will make you a stronger analyst. These details are "
        "what separates a good SOC report from a great one."
    )
    notice_run.font.size      = Pt(10)
    notice_run.font.color.rgb = RGBColor(0, 70, 127)

    doc.add_paragraph()

    # ── Header ────────────────────────────────────────────────────
    title = "STUDENT FEEDBACK REPORT"
    if is_regrade:
        title += f"  ({regrade_ver.upper()})"
    add_heading(title, level=1, color=(0, 70, 127))
    doc.add_paragraph()

    table       = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Student:", student_name),
        ("Lab:",     f"Lab {lab_num}"),
        ("Date:",    datetime.now().strftime("%Y-%m-%d")),
    ]):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────
    add_heading("EXECUTIVE SUMMARY", level=2, color=(0, 70, 127))

    score_pct   = round((score / total) * 100) if total > 0 else 0
    adjusted    = min(score + 20, total)
    score_color = (0, 128, 0) if score_pct >= 70 else (200, 0, 0)

    add_para(f"Raw Score:       {score} / {total}", bold=True, size=12)
    add_para(
        f"Adjusted Score:  {adjusted} / {total}  ({min(score_pct + 20, 100)}%)",
        bold=True, size=13, color=score_color
    )
    doc.add_paragraph()

    if "missing required screenshot" in summary_text.lower():
        add_para(
            "\u26a0  One or more required screenshots were missing. "
            "See detailed feedback. Text answer points are preserved "
            "but your instructor may not accept answers without "
            "screenshots per lab policy.",
            bold=True, size=10, color=(180, 80, 0)
        )
        doc.add_paragraph()

    # ── Question Score Matrix ─────────────────────────────────────
    for line in summary_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("Total Score:") \
                or line.startswith("Points Possible:") \
                or line.startswith("Sum of"):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        if line.startswith("Q:"):
            run = p.add_run(line)
            run.font.color.rgb = (
                RGBColor(0, 128, 0) if "Correct" in line
                else RGBColor(180, 0, 0)
            )
        else:
            p.add_run(line)

    doc.add_paragraph()

    # ── Detailed Feedback ─────────────────────────────────────────
    if detailed_text.strip() and "all questions answered correctly" \
            not in detailed_text.lower():
        add_heading("DETAILED FEEDBACK", level=2, color=(0, 70, 127))
        add_para(
            "The following questions had deductions. "
            "Review each item carefully before resubmitting.",
            size=10
        )
        doc.add_paragraph()

        for block in re.split(r'(?=QUESTION:)', detailed_text):
            block = block.strip()
            if not block:
                continue
            for line in block.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                if line.startswith("QUESTION:"):
                    p = doc.add_heading(
                        line.replace("QUESTION:", "").strip(), level=3
                    )
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(180, 0, 0)
                elif line.startswith("SCREENSHOT_STATUS:"):
                    status = line.replace("SCREENSHOT_STATUS:", "").strip()
                    color  = (180, 80, 0) if "Missing" in status else (100, 100, 100)
                    add_para(f"Screenshot Status:  {status}",
                             bold=True, size=10, color=color, indent=0.25)
                elif line.startswith("WHAT_YOU_SUBMITTED:"):
                    add_para("What You Submitted:", bold=True)
                    add_para(line.replace("WHAT_YOU_SUBMITTED:", "").strip(), indent=0.25)
                elif line.startswith("WHAT_WAS_WRONG:"):
                    add_para("What Was Wrong:", bold=True)
                    add_para(line.replace("WHAT_WAS_WRONG:", "").strip(), indent=0.25)
                elif line.startswith("WHY_IT_MATTERS:"):
                    add_para("Why It Matters:", bold=True)
                    add_para(line.replace("WHY_IT_MATTERS:", "").strip(), indent=0.25)
                elif line.startswith("HOW_TO_FIX:"):
                    add_para("How To Fix It:", bold=True)
                    add_para(line.replace("HOW_TO_FIX:", "").strip(), indent=0.25)
                else:
                    add_para(line)
            doc.add_paragraph("\u2500" * 60)

    doc.save(str(output_path))
    return output_path


def write_grade_log(course_dir: Path, initials: str, student_name: str,
                    lab_num: int, score: int, total: int,
                    feedback_file: str, regrade_ver: str | None,
                    page_count: int, batch_id: str):
    """Append grade record to grade_log.csv."""
    log_file    = course_dir / "grade_log.csv"
    file_exists = log_file.exists()
    with open(log_file, "a", newline="", encoding="utf-8") as f:
        fieldnames = [
            "timestamp", "initials", "student_name", "lab",
            "score", "adjusted_score", "total", "percent",
            "regrade", "agent_version", "pages_rendered",
            "batch_id", "feedback_file"
        ]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        if not file_exists:
            writer.writeheader()
        pct      = round((score / total) * 100) if total > 0 else 0
        adjusted = min(score + 20, total)
        writer.writerow({
            "timestamp":      datetime.now().strftime("%Y-%m-%d %H:%M"),
            "initials":       initials,
            "student_name":   student_name,
            "lab":            lab_num,
            "score":          score,
            "adjusted_score": adjusted,
            "total":          total,
            "percent":        f"{min(pct + 20, 100)}%",
            "regrade":        regrade_ver if regrade_ver else "No",
            "agent_version":  AGENT_VERSION,
            "pages_rendered": page_count,
            "batch_id":       batch_id,
            "feedback_file":  feedback_file,
        })


def main():
    print("\n" + "="*54)
    print(f"  Cisco Cyber Ops  \u2014  Batch Retrieve {AGENT_VERSION}")
    print("="*54)

    course_dir = pick_course_folder()
    print(f"\n  Working in:  {course_dir.name}")

    # Load pending jobs
    pending_jobs = load_batch_manifest(course_dir)
    if not pending_jobs:
        print("\n  No pending batch jobs found in batch_jobs.csv.")
        print("  Run grading_agentV5.py first to submit a batch.\n")
        sys.exit(0)

    # Get unique batch IDs
    batch_ids = list(dict.fromkeys(j["batch_id"] for j in pending_jobs))
    print(f"\n  Found {len(pending_jobs)} pending result(s) across "
          f"{len(batch_ids)} batch job(s).")

    client = anthropic.Anthropic(api_key=API_KEY)
    total_written = 0

    for batch_id in batch_ids:
        print(f"\n  {'='*54}")
        print(f"  Checking batch:  {batch_id}")

        try:
            batch = client.messages.batches.retrieve(batch_id)
        except Exception as e:
            print(f"  ERROR: Could not retrieve batch — {e}")
            continue

        print(f"  Status:  {batch.processing_status}")

        if batch.processing_status != "ended":
            print(f"  Not ready yet.  Try again later.")
            continue

        # Get jobs for this batch
        batch_jobs = {
            j["custom_id"]: j
            for j in pending_jobs
            if j["batch_id"] == batch_id
        }

        # Stream results
        print(f"  Downloading results...")
        try:
            results = client.messages.batches.results(batch_id)
        except Exception as e:
            print(f"  ERROR: Could not download results — {e}")
            continue

        for result in results:
            custom_id = result.custom_id
            job       = batch_jobs.get(custom_id)
            if not job:
                continue

            print(f"\n  Processing:  {custom_id}")

            if result.result.type == "error":
                print(f"  ERROR: API returned error for {custom_id}")
                print(f"  {result.result.error}")
                update_manifest_status(course_dir, custom_id, "error")
                continue

            # Extract response text
            response_text = result.result.message.content[0].text

            # Parse grade
            summary_text, detailed_text, score, total = parse_grade_response(
                response_text
            )

            # Parse job metadata
            initials    = job["initials"]
            student_name = job["student_name"]
            lab_num     = int(job["lab"])
            regrade_ver = job["regrade"] if job["regrade"] != "No" else None
            page_count  = int(job.get("page_count", 0))
            adjusted    = min(score + 20, total)

            print(f"  Student:        {student_name}")
            print(f"  Raw Score:      {score}/{total}")
            print(f"  Adjusted Score: {adjusted}/{total}")

            # Write feedback docx
            feedback_path = write_feedback_docx(
                course_dir, student_name, initials, lab_num,
                summary_text, detailed_text, score, total, regrade_ver
            )
            print(f"  Feedback:       {feedback_path.name}")

            # Log grade
            write_grade_log(
                course_dir, initials, student_name, lab_num,
                score, total, feedback_path.name, regrade_ver,
                page_count, batch_id
            )

            # Update manifest
            update_manifest_status(course_dir, custom_id, "completed")
            total_written += 1
            print(f"  ✓  Complete")

    print(f"\n{'='*54}")
    print(f"  Done.  {total_written} feedback document(s) written.")
    print(f"  Grades logged to grade_log.csv")
    print(f"  Feedback files in grades/lab##/ folders")
    print(f"{'='*54}\n")


if __name__ == "__main__":
    main()
