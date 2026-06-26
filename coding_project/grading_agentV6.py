"""
grading_agentV6.py  —  Automated Grading Agent (Batch Submission)
Cisco Cyber Operations Grading System

Changes from V5:
  - Stronger output format enforcement for claude-sonnet-4-6
  - Two explicit format reminders added to prompt — suppresses narrative thinking
  - Fixes issue where Sonnet returned narrative analysis instead of structured output
  - Fixes raw score returning 0 when EXECUTIVE_SUMMARY_START marker was missing

Usage:
    python grading_agentV6.py    <- submits batch, exits immediately
    python batch_retrieveV1.py   <- next day, collects results

Dependencies:
    pip install anthropic python-dotenv python-docx pdf2image
    poppler at: C:\\poppler\\poppler-26.02.0\\Library\\bin
"""

import os
import sys
import csv
import json
import base64
import re
from datetime import datetime
from pathlib import Path
from io import BytesIO

import anthropic
from dotenv import load_dotenv
from docx import Document
from pdf2image import convert_from_path

# ── Load API key ──────────────────────────────────────────────────────────────
load_dotenv(Path("I:/Visual_Studio_Code/.env"))
API_KEY = os.getenv("ANTHROPIC_API_KEY")
if not API_KEY:
    print("\n  ERROR: ANTHROPIC_API_KEY not found in .env file.")
    print("  Check I:\\Visual_Studio_Code\\.env has your API key.\n")
    sys.exit(1)

# ── Configuration ─────────────────────────────────────────────────────────────
BASE_DIR      = Path("I:/Visual_Studio_Code")
MODEL         = "claude-sonnet-4-6"
MAX_TOKENS    = 8192
AGENT_VERSION = "V6"
POPPLER_PATH  = r"C:\poppler\poppler-26.02.0\Library\bin"
MIN_PAGES     = 3
RENDER_DPI    = 100    # reduced from 150 to keep batch under 256MB limit
# ──────────────────────────────────────────────────────────────────────────────


def pick_course_folder() -> Path:
    course_folders = sorted([
        f for f in BASE_DIR.iterdir()
        if f.is_dir() and (f / "roster.csv").exists()
    ])
    if not course_folders:
        print(f"\n  ERROR: No course folders found under {BASE_DIR}")
        sys.exit(1)
    print("\n  ── Course Folders Found ────────────────────────────")
    for i, folder in enumerate(course_folders, start=1):
        print(f"  {i:>3}.  {folder.name}")
    print("  ────────────────────────────────────────────────────")
    while True:
        raw = input("\n  Enter course number: ").strip()
        if raw.isdigit():
            num = int(raw)
            if 1 <= num <= len(course_folders):
                selected = course_folders[num - 1]
                confirm  = input(f"  Confirm: {selected.name}  (y/n): ").strip().lower()
                if confirm == "y":
                    return selected
                print("  Not confirmed.  Please re-enter.")
                continue
        print(f"  Invalid.  Enter a number between 1 and {len(course_folders)}.")


def load_roster(course_dir: Path) -> dict:
    roster_file = course_dir / "roster.csv"
    lookup = {}
    with open(roster_file, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        reader.fieldnames = [h.strip().lower() for h in reader.fieldnames]
        for row in reader:
            clean    = {k.strip().lower(): v for k, v in row.items()}
            initials = clean["initials"].strip().upper()
            lookup[initials] = (
                f"{clean['first_name'].strip()} {clean['last_name'].strip()}"
            )
    return lookup


def parse_queue_filename(filename: str):
    """Extract initials, lab number, regrade version from filename."""
    stem          = Path(filename).stem
    regrade_match = re.search(r'_r(\d+)$', stem)
    regrade_ver   = f"r{regrade_match.group(1)}" if regrade_match else None
    match         = re.match(r'^([A-Z0-9]+)_lab(\d+)_pending', stem, re.IGNORECASE)
    if match:
        return match.group(1).upper(), int(match.group(2)), regrade_ver
    return None, None, None


def find_template_files(course_dir: Path, lab_num: int):
    templates_dir = course_dir / "templates"
    answers_file  = templates_dir / f"2026SU_Lab_{lab_num}_-_ILM.docx"
    rubric_file   = templates_dir / f"Lab_{lab_num}_Grading_Notes.txt"
    answers       = answers_file if answers_file.exists() else None
    rubric        = rubric_file  if rubric_file.exists()  else None
    return answers, rubric


def extract_template_answers(file_path: Path) -> str:
    """
    Extract ONLY answer-relevant content from the template docx.
    Strips lab instructions, background text, and click-by-click steps.
    Keeps: correct answers, screenshot markers, NOT-used statements,
           question labels, question text, answer tables.
    Reduces template tokens by approximately 89%.
    """
    try:
        doc   = Document(file_path)
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        return f"[Could not extract template: {e}]"

    output = []
    for line in lines:
        lower = line.lower()

        # Always keep — table rows handled separately via python-docx
        # Keep screenshot requirement markers
        if 'insert your screenshot' in lower:
            output.append(line)
            continue

        # Keep the NOT used to answer statement — critical for grading logic
        if 'not' in lower and 'used' in lower and 'answer' in lower:
            output.append(line)
            continue

        # Keep answer insert markers
        if 'insert your answer' in lower:
            output.append(line)
            continue

        # Keep Part/Step labels
        if re.match(r'^Part \d+', line) or re.match(r'^Step \d+', line):
            output.append(line)
            continue

        # Keep question text — lines ending with ?
        if line.endswith('?') and len(line) < 300:
            output.append(line)
            continue

        # Keep short bold answer lines (correct answers in template)
        # These are wrapped in ** in extracted markdown
        if line.startswith('**') and len(line) < 150:
            output.append(line)
            continue

        # Keep lettered sub-questions like "a." "b." etc
        if re.match(r'^[a-z]\.\s', line) and len(line) < 300:
            output.append(line)
            continue

        # Keep numbered sub-questions
        if re.match(r'^\d+\.\s', line) and len(line) < 300:
            output.append(line)
            continue

    # Also extract table content from docx tables directly
    try:
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    output.append(" | ".join(cells))
    except Exception:
        pass

    return "\n".join(output)


def read_text_file(file_path: Path) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def pdf_to_images(pdf_path: Path) -> list:
    return convert_from_path(
        str(pdf_path),
        dpi=RENDER_DPI,
        poppler_path=POPPLER_PATH,
        fmt="jpeg",
        thread_count=2,
    )


def image_to_base64(pil_image) -> str:
    buffer = BytesIO()
    pil_image.save(buffer, format="JPEG", quality=70)   # reduced from 85 to keep batch under 256MB limit
    return base64.standard_b64encode(buffer.getvalue()).decode("utf-8")


def verify_and_render_pdf(pdf_path: Path, failed_dir: Path):
    """Verify PDF renders and return pages. Returns (success, pages)."""
    print(f"  Verifying PDF...")
    try:
        pages = pdf_to_images(pdf_path)
    except Exception as e:
        print(f"  ERROR: PDF conversion failed — {e}")
        failed_dir.mkdir(exist_ok=True)
        import shutil
        shutil.move(str(pdf_path),
                    str(failed_dir / f"{pdf_path.stem}_conversion_failed{pdf_path.suffix}"))
        return False, []

    if len(pages) < MIN_PAGES:
        print(f"  ERROR: Only {len(pages)} page(s) — too few.")
        failed_dir.mkdir(exist_ok=True)
        import shutil
        shutil.move(str(pdf_path),
                    str(failed_dir / f"{pdf_path.stem}_too_few_pages{pdf_path.suffix}"))
        return False, []

    print(f"  ✓  {len(pages)} pages ready.")
    return True, pages


def build_grading_prompt(answers_text: str, rubric_text: str,
                         student_name: str, lab_num: int,
                         page_count: int) -> str:
    return f"""CRITICAL INSTRUCTION: Your response must begin immediately with EXECUTIVE_SUMMARY_START.
Do not write any analysis, reasoning, narrative, or thinking text before EXECUTIVE_SUMMARY_START.
Structured output only. Start with EXECUTIVE_SUMMARY_START on the very first line.

You are an experienced Cisco Cyber Operations instructor grading a student lab submission.

STUDENT: {student_name}
LAB: {lab_num}

You have been provided:
1. ANSWER TEMPLATE — correct answers extracted from the lab template
2. GRADING NOTES — lead instructor rubric with specific deductions
3. STUDENT SUBMISSION — {page_count} pages rendered as images

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
CRITICAL \u2014 OUTPUT FORMAT REQUIREMENTS
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

You MUST respond using ONLY the structured format below.
Do NOT write a narrative review.
Do NOT describe what you are doing.
Start your response immediately with EXECUTIVE_SUMMARY_START.
Every question graded must appear as a Q: line in the matrix.

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
SCREENSHOT GRADING RULES
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

TYPE 1 \u2014 SCREENSHOT IS THE ANSWER
  Template asks for screenshot with no qualifying statement below it.
  Grade screenshot directly. If missing or incorrect deduct full points.
  If present but unclear deduct partial points and explain.

TYPE 2 \u2014 SCREENSHOT REQUIRED BUT NOT THE ANSWER
  Template contains: "Your screenshot is NOT used to answer the questions below"
  STEP A: Verify screenshot exists on previous page (pass/fail only)
    If MISSING: flag as separate deduction, still grade text answers,
    warn instructor may not accept answers without screenshot.
    If PRESENT: do NOT use screenshot content to grade text answers.
  STEP B: Grade TEXT ANSWERS below the statement independently.

SCREENSHOT FORMATTING DEDUCTIONS (max 11 points total, separate from questions):
  - Too small to read output clearly
  - Wrong area of screen captured
  - NDG username not visible
  - Multiple screenshots where one was required

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
SCORING RULES
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

- Grade every question individually first
- Each question score stated as [earned]/[possible]
- Points per question is 6 unless grading notes state otherwise
- Calculate Total Score by SUMMING every Q: line — do NOT estimate
- Screenshot formatting deductions are separate from question scores
- Use grading notes as PRIMARY source — expand language, do not copy
- Connect errors to real Cisco Cyber Ops or SOC security concepts

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
REQUIRED OUTPUT \u2014 START WITH EXECUTIVE_SUMMARY_START
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

EXECUTIVE_SUMMARY_START
Points Possible: [total points in this lab]
Screenshot Formatting Deduction: [X] points
Missing Required Screenshots: [list question labels or None]

QUESTION SCORE MATRIX:
Q: [Part X Step Xx] | [earned]/[possible] | Correct
Q: [Part X Step Xx] | [earned]/[possible] | [brief reason for deduction]
[one Q: line for every graded question \u2014 do not skip any]

SCORE CALCULATION:
Sum of earned points: [add every earned score from Q: lines above]
Screenshot deduction: [X]
Total Score: [sum minus screenshot deduction]/[Points Possible]
EXECUTIVE_SUMMARY_END

DETAILED_FEEDBACK_START
[Only questions where points were deducted]
[If all correct write: All questions answered correctly.]

QUESTION: [Part X Step Xx] \u2014 [earned]/[possible] points
SCREENSHOT_STATUS: [Not applicable / Present / Missing / Present but unclear]
WHAT_YOU_SUBMITTED: [brief summary of what student wrote or showed]
WHAT_WAS_WRONG: [specific explanation of the error]
WHY_IT_MATTERS: [connect to Cisco Cyber Ops or SOC security concept]
HOW_TO_FIX: [clear corrective guidance for resubmission]

[repeat for each incorrect question only]
DETAILED_FEEDBACK_END

\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550
FINAL REMINDER BEFORE YOU BEGIN


Your first line of output must be: EXECUTIVE_SUMMARY_START
No exceptions. No narrative before it. No thinking out loud.
Grade every question. Sum the Q: matrix. Then write EXECUTIVE_SUMMARY_END.
Then write DETAILED_FEEDBACK_START through DETAILED_FEEDBACK_END.
Nothing else.

GRADING MATERIALS
\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550\u2550

ANSWER TEMPLATE (answers and screenshot requirements only):
{answers_text}

GRADING NOTES:
{rubric_text}
"""


def build_batch_request(custom_id: str, prompt: str,
                        pages: list, file: Path) -> dict:
    """Build one batch request object for a student submission."""
    ext = file.suffix.lower()

    if ext == ".pdf" and pages:
        content = []
        for i, page in enumerate(pages):
            content.append({
                "type": "image",
                "source": {
                    "type":       "base64",
                    "media_type": "image/jpeg",
                    "data":       image_to_base64(page),
                }
            })
            content.append({
                "type": "text",
                "text": f"[Page {i + 1} of {len(pages)}]"
            })
        content.append({
            "type": "text",
            "text": (
                "The pages above are the complete student submission. "
                "Grade this submission per the rubric and answer template "
                "in your instructions."
            )
        })
    else:
        # DOCX fallback
        from docx import Document as DocxDoc
        try:
            doc  = DocxDoc(file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            text = "[Could not extract docx text]"
        content = (
            "Grade this student submission per the rubric and answer template "
            f"in your instructions.\n\nSTUDENT SUBMISSION:\n{text}"
        )

    return {
        "custom_id": custom_id,
        "params": {
            "model":      MODEL,
            "max_tokens": MAX_TOKENS,
            "system":     prompt,
            "messages":   [{"role": "user", "content": content}],
        }
    }


def save_batch_manifest(course_dir: Path, batch_id: str,
                        submissions: list[dict]):
    """
    Save batch job manifest to batch_jobs.csv so retrieve script
    knows which files belong to this batch.
    """
    manifest_file = course_dir / "batch_jobs.csv"
    file_exists   = manifest_file.exists()
    with open(manifest_file, "a", newline="", encoding="utf-8") as f:
        fieldnames = [
            "batch_id", "submitted_at", "custom_id",
            "initials", "student_name", "lab", "regrade",
            "original_file", "page_count", "status"
        ]
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        if not file_exists:
            writer.writeheader()
        for s in submissions:
            writer.writerow({
                "batch_id":      batch_id,
                "submitted_at":  datetime.now().strftime("%Y-%m-%d %H:%M"),
                "custom_id":     s["custom_id"],
                "initials":      s["initials"],
                "student_name":  s["student_name"],
                "lab":           s["lab"],
                "regrade":       s["regrade"] or "No",
                "original_file": s["original_file"],
                "page_count":    s["page_count"],
                "status":        "pending",
            })


def main():
    print("\n" + "="*54)
    print(f"  Cisco Cyber Ops  \u2014  Grading Agent {AGENT_VERSION}")
    print(f"  Batch Submission Mode")
    print("="*54)

    course_dir = pick_course_folder()
    print(f"\n  Working in:  {course_dir.name}")

    roster    = load_roster(course_dir)
    client    = anthropic.Anthropic(api_key=API_KEY)
    queue_dir = course_dir / "queue"
    failed_dir = course_dir / "failed"

    # Find all pending files
    pending = sorted([
        f for f in queue_dir.iterdir()
        if f.is_file()
        and "_pending" in f.name.lower()
        and f.suffix.lower() in {".pdf", ".docx"}
    ])

    if not pending:
        print("\n  No pending files found in queue folder.")
        print("  Run intakeV2.py first to add student submissions.\n")
        sys.exit(0)

    print(f"\n  Found {len(pending)} file(s) in queue:")
    for f in pending:
        print(f"    {f.name}")

    confirm = input(f"\n  Submit all {len(pending)} files as batch job? (y/n): ").strip().lower()
    if confirm != "y":
        print("  Cancelled.\n")
        sys.exit(0)

    # Build batch requests
    batch_requests  = []
    submissions_log = []
    skipped         = 0

    for file in pending:
        print(f"\n  Preparing:  {file.name}")

        initials, lab_num, regrade_ver = parse_queue_filename(file.name)
        if not initials or not lab_num:
            print(f"  ERROR: Cannot parse filename — skipping.")
            skipped += 1
            continue

        student_name = roster.get(initials, initials)
        print(f"  Student:    {student_name}")
        print(f"  Lab:        {lab_num}")

        answers_path, rubric_path = find_template_files(course_dir, lab_num)
        if not answers_path or not rubric_path:
            print(f"  ERROR: Template or rubric not found for Lab {lab_num} — skipping.")
            skipped += 1
            continue

        # Render PDF pages
        pages      = []
        page_count = 0
        if file.suffix.lower() == ".pdf":
            ok, pages = verify_and_render_pdf(file, failed_dir)
            if not ok:
                skipped += 1
                continue
            page_count = len(pages)

        # Extract template answers only — 89% token reduction
        answers_text = extract_template_answers(answers_path)
        rubric_text  = read_text_file(rubric_path)
        prompt       = build_grading_prompt(
            answers_text, rubric_text, student_name, lab_num, page_count
        )

        # Build unique custom_id for this request
        ver_suffix = f"_{regrade_ver}" if regrade_ver else ""
        custom_id  = f"{initials}_lab{lab_num:02d}{ver_suffix}"

        batch_requests.append(
            build_batch_request(custom_id, prompt, pages, file)
        )

        submissions_log.append({
            "custom_id":     custom_id,
            "initials":      initials,
            "student_name":  student_name,
            "lab":           lab_num,
            "regrade":       regrade_ver,
            "original_file": file.name,
            "page_count":    page_count,
        })

        print(f"  ✓  Request prepared  ({page_count} pages)")

    if not batch_requests:
        print("\n  No valid requests to submit.\n")
        sys.exit(0)

    # Submit batch to Anthropic
    print(f"\n  {'='*54}")
    print(f"  Submitting {len(batch_requests)} request(s) to Batch API...")
    print(f"  {'='*54}")

    try:
        batch = client.messages.batches.create(requests=batch_requests)
    except Exception as e:
        print(f"\n  ERROR: Batch submission failed — {e}\n")
        sys.exit(1)

    batch_id = batch.id
    print(f"\n  ✓  Batch submitted successfully.")
    print(f"  Batch ID:  {batch_id}")
    print(f"  Status:    {batch.processing_status}")

    # Save manifest
    save_batch_manifest(course_dir, batch_id, submissions_log)

    # Move submitted files to processed folder
    import shutil
    processed_dir = course_dir / "processed"
    processed_dir.mkdir(exist_ok=True)
    for file in pending:
        if any(s["original_file"] == file.name for s in submissions_log):
            dest = processed_dir / file.name
            if dest.exists():
                ts   = datetime.now().strftime("%H%M%S")
                dest = processed_dir / f"{file.stem}_{ts}{file.suffix}"
            shutil.move(str(file), str(dest))

    print(f"\n  Submissions moved to processed/ folder.")
    print(f"\n  {'='*54}")
    print(f"  Batch job submitted.  Results available within 24 hours.")
    print(f"  Run batch_retrieveV1.py to collect results and write feedback.")
    print(f"  {'='*54}\n")


if __name__ == "__main__":
    main()
